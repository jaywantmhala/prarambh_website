import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_report():
    doc = Document()

    # Define Colors
    COLOR_PRIMARY = RGBColor(43, 58, 66)      # Slate Blue: #2B3A42
    COLOR_SECONDARY = RGBColor(255, 111, 0)   # Solar Orange: #FF6F00
    COLOR_TEXT = RGBColor(51, 51, 51)         # Charcoal: #333333
    COLOR_MUTED = RGBColor(115, 115, 115)     # Muted Gray: #737373
    COLOR_GREEN = RGBColor(22, 163, 74)       # Success Green: #16A34A
    
    HEX_SECONDARY = 'FF6F00'
    HEX_GREEN = '16A34A'
    HEX_BG_LIGHT_ORANGE = 'FFF8F2'            # Shading for Good Video Callout
    HEX_BG_LIGHT_GREEN = 'F0FDF4'             # Shading for Improvements Callout

    # Set margins to 1 inch for all sections
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Setup normal body text style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.25
    style_normal.paragraph_format.space_after = Pt(6)

    # Helper function to style text runs
    def format_run(run, font_name='Segoe UI', size_pt=11, bold=False, italic=False, color=COLOR_TEXT):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color

    # XML Helper: Shading a table cell
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

    # XML Helper: Cell margins/padding (in dxa)
    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # XML Helper: Custom thick left border (for callout boxes) and remove others
    def set_cell_left_border(cell, border_hex, size=32):  # 32/8 = 4pt
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        # Left border
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(size))
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), border_hex)
        tcBorders.append(left)
        
        # Top, Bottom, Right borders set to none
        for side in ['top', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
            
        tcPr.append(tcBorders)

    # XML Helper: Center text vertically in a section
    def set_section_vertical_alignment(section, alignment="center"):
        sectPr = section._sectPr
        # Remove any existing w:vAlign tag to prevent duplicates
        for el in sectPr:
            if el.tag.endswith('vAlign'):
                sectPr.remove(el)
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), alignment)
        sectPr.append(vAlign)

    # Helper to add clickable hyperlinks in docx
    def add_hyperlink(paragraph, url, text, color="FF6F00", underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)

        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        new_run.append(rPr)

        text_node = OxmlElement('w:t')
        text_node.text = text
        new_run.append(text_node)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    # Helper: Add beautiful callout box
    def add_callout_box(doc, title, text, bg_hex, border_hex):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        
        cell = table.cell(0, 0)
        set_cell_background(cell, bg_hex)
        set_cell_left_border(cell, border_hex, size=32) # 4pt Left Border
        set_cell_margins(cell, top=160, bottom=160, left=200, right=160)
        
        # Add callout title
        p_title = cell.paragraphs[0]
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run(title)
        format_run(run_title, size_pt=11, bold=True, color=COLOR_PRIMARY if border_hex == HEX_SECONDARY else COLOR_GREEN)
        
        # Add callout description
        p_desc = cell.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(0)
        p_desc.paragraph_format.line_spacing = 1.2
        run_desc = p_desc.add_run(text)
        format_run(run_desc, size_pt=10.5, italic=True)

    # ----------------------------------------------------
    # PAGE 1: CENTERED TITLE & COMPANY DETAILS
    # ----------------------------------------------------
    # Apply vertical centering to first section
    first_section = doc.sections[0]
    set_section_vertical_alignment(first_section, "center")

    # Document Header Title
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_after = Pt(4)
    r_main_title = p_main_title.add_run("VIDEO ANALYSIS REPORT")
    format_run(r_main_title, size_pt=26, bold=True, color=COLOR_PRIMARY)

    # Subtitle / Company Name
    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_comp.paragraph_format.space_after = Pt(36)
    r_comp = p_comp.add_run("AYKA SOLAR")
    format_run(r_comp, size_pt=18, bold=True, color=COLOR_SECONDARY)

    # Website Label
    p_web_lbl = doc.add_paragraph()
    p_web_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_web_lbl.paragraph_format.space_after = Pt(2)
    r_web_lbl = p_web_lbl.add_run("Company Website:")
    format_run(r_web_lbl, size_pt=10.5, bold=True, color=COLOR_MUTED)

    # Website Clickable Hyperlink
    p_web_lnk = doc.add_paragraph()
    p_web_lnk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_web_lnk.paragraph_format.space_after = Pt(24)
    add_hyperlink(p_web_lnk, "https://ayka.com.au/?utm_source=chatgpt.com", "ayka.com.au")

    # Instagram Label
    p_ig_lbl = doc.add_paragraph()
    p_ig_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ig_lbl.paragraph_format.space_after = Pt(2)
    r_ig_lbl = p_ig_lbl.add_run("Instagram Profile:")
    format_run(r_ig_lbl, size_pt=10.5, bold=True, color=COLOR_MUTED)

    # Instagram Clickable Hyperlink
    p_ig_lnk = doc.add_paragraph()
    p_ig_lnk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ig_lnk.paragraph_format.space_after = Pt(0)
    add_hyperlink(p_ig_lnk, "https://www.instagram.com/aykasolar_australia/?utm_source=chatgpt.com", "@aykasolar_australia")

    # ----------------------------------------------------
    # PAGE 2: GOOD VIDEO 1 (Animated Educational Reel)
    # ----------------------------------------------------
    # Add a new section for top-aligned subsequent pages
    second_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_vertical_alignment(second_section, "top")
    
    # Ensure margins match
    second_section.top_margin = Inches(1)
    second_section.bottom_margin = Inches(1)
    second_section.left_margin = Inches(1)
    second_section.right_margin = Inches(1)

    p_g1_title = doc.add_paragraph()
    p_g1_title.paragraph_format.space_after = Pt(2)
    format_run(p_g1_title.add_run("GOOD VIDEO 1"), size_pt=18, bold=True, color=COLOR_PRIMARY)

    p_g1_sub = doc.add_paragraph()
    p_g1_sub.paragraph_format.space_after = Pt(12)
    format_run(p_g1_sub.add_run("Animated Educational Reel"), size_pt=12, italic=True, color=COLOR_SECONDARY)

    # Video Link
    p_g1_lnk = doc.add_paragraph()
    p_g1_lnk.paragraph_format.space_after = Pt(6)
    r_g1_lnk_lbl = p_g1_lnk.add_run("Video Link: ")
    format_run(r_g1_lnk_lbl, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
    add_hyperlink(p_g1_lnk, "https://www.instagram.com/reel/DXyat1wkomI/?igsh=cDlkbG55NDAzajc3", "https://www.instagram.com/reel/DXyat1wkomI/")

    # Type
    p_g1_type = doc.add_paragraph()
    p_g1_type.paragraph_format.space_after = Pt(16)
    r_g1_type_lbl = p_g1_type.add_run("Type: ")
    format_run(r_g1_type_lbl, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
    r_g1_type_val = p_g1_type.add_run("Animated educational reel about underperforming solar systems")
    format_run(r_g1_type_val, size_pt=10.5, color=COLOR_TEXT)

    # Why Good
    p_g1_why_lbl = doc.add_paragraph()
    p_g1_why_lbl.paragraph_format.space_after = Pt(6)
    format_run(p_g1_why_lbl.add_run("Why this video is good:"), size_pt=12, bold=True, color=COLOR_PRIMARY)

    bullets_g1 = [
        "The reel starts with a question hook: “Is your system actually working?” which creates curiosity.",
        "Animated visuals make technical information easy to understand.",
        "Bright colors and clean graphics grab attention quickly.",
        "The content is educational and valuable for homeowners.",
        "Subtitles improve watch time and accessibility.",
        "The reel explains a problem and creates awareness which helps audience engagement.",
        "Clear branding and contact details improve trust and lead generation."
    ]

    for b in bullets_g1:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        
        text = p_b.add_run(b)
        format_run(text, size_pt=10.5, color=COLOR_TEXT)

    doc.add_paragraph() # Spacing

    # What Works Box
    add_callout_box(
        doc,
        "What works:",
        "This reel works well because it mixes education with marketing. Instead of directly selling, it first informs viewers about a possible problem in their solar system, making the audience more interested in the service.",
        HEX_BG_LIGHT_ORANGE,
        HEX_SECONDARY
    )

    # ----------------------------------------------------
    # PAGE 3: GOOD VIDEO 2 (Funny Relatable Meme)
    # ----------------------------------------------------
    doc.add_page_break()

    p_g2_title = doc.add_paragraph()
    p_g2_title.paragraph_format.space_after = Pt(2)
    format_run(p_g2_title.add_run("GOOD VIDEO 2"), size_pt=18, bold=True, color=COLOR_PRIMARY)

    p_g2_sub = doc.add_paragraph()
    p_g2_sub.paragraph_format.space_after = Pt(12)
    format_run(p_g2_sub.add_run("Funny Relatable Meme-Style Reel"), size_pt=12, italic=True, color=COLOR_SECONDARY)

    # Video Link
    p_g2_lnk = doc.add_paragraph()
    p_g2_lnk.paragraph_format.space_after = Pt(6)
    r_g2_lnk_lbl = p_g2_lnk.add_run("Video Link: ")
    format_run(r_g2_lnk_lbl, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
    add_hyperlink(p_g2_lnk, "https://www.instagram.com/reel/DQqqNf4jds3/?igsh=MTJ6YXN4eTk2YnI3Mg==", "https://www.instagram.com/reel/DQqqNf4jds3/")

    # Type
    p_g2_type = doc.add_paragraph()
    p_g2_type.paragraph_format.space_after = Pt(16)
    r_g2_type_lbl = p_g2_type.add_run("Type: ")
    format_run(r_g2_type_lbl, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
    r_g2_type_val = p_g2_type.add_run("Funny relatable meme-style reel")
    format_run(r_g2_type_val, size_pt=10.5, color=COLOR_TEXT)

    # Why Good
    p_g2_why_lbl = doc.add_paragraph()
    p_g2_why_lbl.paragraph_format.space_after = Pt(6)
    format_run(p_g2_why_lbl.add_run("Why this video is good:"), size_pt=12, bold=True, color=COLOR_PRIMARY)

    bullets_g2 = [
        "The reel uses a popular Tom & Jerry meme format which instantly grabs attention.",
        "Funny and relatable content increases shareability.",
        "Simple message is easy to understand within seconds.",
        "Humor makes the brand feel more modern and engaging.",
        "Meme-based content performs well on Instagram because audiences enjoy entertaining reels.",
        "The text “After installing solar system with Battery” clearly communicates the benefit in a fun way.",
        "Short-form meme content improves audience retention."
    ]

    for b in bullets_g2:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        
        text = p_b.add_run(b)
        format_run(text, size_pt=10.5, color=COLOR_TEXT)

    doc.add_paragraph() # Spacing

    # What Works Box
    add_callout_box(
        doc,
        "What works:",
        "This reel works because it combines humor with marketing. Instead of directly selling the product, it shows the comfort and convenience of having a solar battery system in a relatable and entertaining way.",
        HEX_BG_LIGHT_ORANGE,
        HEX_SECONDARY
    )

    # ----------------------------------------------------
    # PAGE 4: NOT GOOD VIDEO 1 (Static Informational Post)
    # ----------------------------------------------------
    doc.add_page_break()

    p_ng1_title = doc.add_paragraph()
    p_ng1_title.paragraph_format.space_after = Pt(2)
    format_run(p_ng1_title.add_run("NOT GOOD VIDEO 1"), size_pt=18, bold=True, color=COLOR_PRIMARY)

    p_ng1_sub = doc.add_paragraph()
    p_ng1_sub.paragraph_format.space_after = Pt(12)
    format_run(p_ng1_sub.add_run("Informational / Static Post"), size_pt=12, italic=True, color=COLOR_MUTED)

    # Type
    p_ng1_type = doc.add_paragraph()
    p_ng1_type.paragraph_format.space_after = Pt(16)
    r_ng1_type_lbl = p_ng1_type.add_run("Type: ")
    format_run(r_ng1_type_lbl, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
    r_ng1_type_val = p_ng1_type.add_run("Informational/static post")
    format_run(r_ng1_type_val, size_pt=10.5, color=COLOR_TEXT)

    # Why Not Good
    p_ng1_why_lbl = doc.add_paragraph()
    p_ng1_why_lbl.paragraph_format.space_after = Pt(6)
    format_run(p_ng1_why_lbl.add_run("Why this video is not good:"), size_pt=12, bold=True, color=COLOR_PRIMARY)

    bullets_ng1 = [
        "Hook is weak and fails to capture immediate focus.",
        "Visuals feel less dynamic compared to video posts.",
        "Text-heavy content may reduce reading retention.",
        "Lacks structural storytelling to keep users engaged.",
        "Audience is highly likely to scroll past quickly."
    ]

    for b in bullets_ng1:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        
        text = p_b.add_run(b)
        format_run(text, size_pt=10.5, color=COLOR_TEXT)

    doc.add_paragraph() # Spacing

    # How to Improve Box
    add_callout_box(
        doc,
        "How to improve:",
        "• Add motion graphics to bring text to life.\n"
        "• Use faster transitions to create rhythm.\n"
        "• Add engaging captions and subtitles.\n"
        "• Start with a stronger hook like “Save electricity bills by 70%”",
        HEX_BG_LIGHT_GREEN,
        HEX_GREEN
    )

    # ----------------------------------------------------
    # PAGE 5: NOT GOOD VIDEO 2 (Simple Promotional Reel)
    # ----------------------------------------------------
    doc.add_page_break()

    p_ng2_title = doc.add_paragraph()
    p_ng2_title.paragraph_format.space_after = Pt(2)
    format_run(p_ng2_title.add_run("NOT GOOD VIDEO 2"), size_pt=18, bold=True, color=COLOR_PRIMARY)

    p_ng2_sub = doc.add_paragraph()
    p_ng2_sub.paragraph_format.space_after = Pt(12)
    format_run(p_ng2_sub.add_run("Simple Service Promotional Reel"), size_pt=12, italic=True, color=COLOR_MUTED)

    # Type
    p_ng2_type = doc.add_paragraph()
    p_ng2_type.paragraph_format.space_after = Pt(16)
    r_ng2_type_lbl = p_ng2_type.add_run("Type: ")
    format_run(r_ng2_type_lbl, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
    r_ng2_type_val = p_ng2_type.add_run("Simple service promotional reel")
    format_run(r_ng2_type_val, size_pt=10.5, color=COLOR_TEXT)

    # Why Not Good
    p_ng2_why_lbl = doc.add_paragraph()
    p_ng2_why_lbl.paragraph_format.space_after = Pt(6)
    format_run(p_ng2_why_lbl.add_run("Why this video is not good:"), size_pt=12, bold=True, color=COLOR_PRIMARY)

    bullets_ng2 = [
        "Editing pace is slow and lacks momentum.",
        "No emotional angle is utilized to connect with the audience.",
        "Background music is not very engaging or modern.",
        "Reel feels more promotional and sales-driven than relatable."
    ]

    for b in bullets_ng2:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        
        text = p_b.add_run(b)
        format_run(text, size_pt=10.5, color=COLOR_TEXT)

    doc.add_paragraph() # Spacing

    # How to Improve Box
    add_callout_box(
        doc,
        "How to improve:",
        "• Add customer problem-solution storytelling elements.\n"
        "• Use clear before/after result comparisons.\n"
        "• Add trending audio to boost platform reach.\n"
        "• Include subtitles and bold text overlays for silent viewing.",
        HEX_BG_LIGHT_GREEN,
        HEX_GREEN
    )

    # ----------------------------------------------------
    # PAGE 6: OVERALL ANALYSIS
    # ----------------------------------------------------
    doc.add_page_break()

    p_oa_title = doc.add_paragraph()
    p_oa_title.paragraph_format.space_after = Pt(14)
    format_run(p_oa_title.add_run("OVERALL ANALYSIS"), size_pt=18, bold=True, color=COLOR_PRIMARY)

    # Strengths Section
    p_str_lbl = doc.add_paragraph()
    p_str_lbl.paragraph_format.space_after = Pt(6)
    format_run(p_str_lbl.add_run("Strengths:"), size_pt=12, bold=True, color=COLOR_PRIMARY)

    strengths = [
        "Professional niche content built around actual industry expertise.",
        "Visuals featuring real quality project installations.",
        "Trust-building, educational format.",
        "Consistent and recognizable branding guidelines."
    ]

    for s in strengths:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        
        text = p_b.add_run(s)
        format_run(text, size_pt=10.5, color=COLOR_TEXT)

    doc.add_paragraph() # Spacing

    # Areas for Improvement Section
    p_imp_lbl = doc.add_paragraph()
    p_imp_lbl.paragraph_format.space_after = Pt(6)
    format_run(p_imp_lbl.add_run("Areas for Improvement:"), size_pt=12, bold=True, color=COLOR_PRIMARY)

    improvements = [
        "Create more engaging, curiosity-inducing hooks (first 3 seconds).",
        "Leverage better narrative storytelling and human-focused angles.",
        "Adopt a faster and snappier video editing style.",
        "Deliver more audience-focused, problem-solving content instead of promotional sales pitches."
    ]

    for i in improvements:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        
        text = p_b.add_run(i)
        format_run(text, size_pt=10.5, color=COLOR_TEXT)

    # Save Document
    doc_path = "AYKA_Solar_Video_Analysis_Report.docx"
    doc.save(doc_path)
    print(f"Document successfully created and saved to: {doc_path}")

if __name__ == "__main__":
    create_report()
